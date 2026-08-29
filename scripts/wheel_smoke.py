#!/usr/bin/env python3
"""Install a built wheel outside the checkout and exercise every CLI path."""

from __future__ import annotations

import argparse
import json
import os
from pathlib import Path
import subprocess
import sys
import tempfile


def _run(command: list[str], *, cwd: Path, env: dict[str, str]) -> None:
    subprocess.run(command, cwd=cwd, env=env, check=True)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("wheel", type=Path)
    parser.add_argument(
        "--system-site-packages",
        action="store_true",
        help="Reuse already installed dependencies for an offline local smoke run.",
    )
    args = parser.parse_args()
    wheel = args.wheel.resolve(strict=True)

    with tempfile.TemporaryDirectory(prefix="er-wheel-smoke-") as temporary:
        root = Path(temporary)
        environment = os.environ.copy()
        environment.pop("PYTHONPATH", None)
        venv_command = [sys.executable, "-m", "venv"]
        if args.system_site_packages:
            venv_command.append("--system-site-packages")
        venv_command.append(str(root / "venv"))
        _run(venv_command, cwd=root, env=environment)

        bin_dir = root / "venv" / ("Scripts" if os.name == "nt" else "bin")
        python = bin_dir / ("python.exe" if os.name == "nt" else "python")
        er_format = bin_dir / ("er-format.exe" if os.name == "nt" else "er-format")
        try:
            _run([str(python), "-m", "pip", "--version"], cwd=root, env=environment)
        except subprocess.CalledProcessError:
            _run([str(python), "-m", "ensurepip", "--upgrade"], cwd=root, env=environment)

        install = [str(python), "-m", "pip", "install"]
        if args.system_site_packages:
            install.append("--no-deps")
        install.append(str(wheel))
        _run(install, cwd=root, env=environment)

        _run([str(er_format), "validate-rules"], cwd=root, env=environment)
        _run(
            [
                str(python),
                "-c",
                (
                    "from docx import Document; "
                    "d=Document(); d.add_paragraph('题目：wheel smoke'); "
                    "d.add_paragraph('摘要'); d.add_paragraph('测试摘要。'); "
                    "d.add_paragraph('关键词：测试'); d.save('minimal.docx')"
                ),
            ],
            cwd=root,
            env=environment,
        )
        _run(
            [str(er_format), "inspect", "minimal.docx", "--output", "inspection.json"],
            cwd=root,
            env=environment,
        )
        _run(
            [
                str(er_format),
                "lint",
                "minimal.docx",
                "--output-dir",
                "audit",
                "--exit-zero",
            ],
            cwd=root,
            env=environment,
        )

        inspection = json.loads((root / "inspection.json").read_text(encoding="utf-8"))
        audit = json.loads((root / "audit" / "audit.json").read_text(encoding="utf-8"))
        assert inspection["schema_version"] == "1.0"
        assert audit["schema_version"] == "1.0"
        assert (root / "audit" / "inspection.json").is_file()
        assert (root / "audit" / "audit.md").is_file()
        location = subprocess.check_output(
            [
                str(python),
                "-c",
                "import economic_research_formatter as p; print(p.__file__)",
            ],
            cwd=root,
            env=environment,
            text=True,
        ).strip()
        assert str(root / "venv") in location
        print(f"wheel_smoke=PASS package={location}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

from __future__ import annotations

import json
import warnings
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document

import economic_research_formatter.cli as cli
import economic_research_formatter.docx.package as package_module
from economic_research_formatter.docx.package import DocxPackage
from economic_research_formatter.models.inspection import DocxInspectionError


W_NS = package_module.W_NS
REL_NS = package_module.REL_NS
CT_NS = package_module.CT_NS


def _minimal_docx(path: Path) -> Path:
    document = Document()
    document.add_paragraph("题目：CLI 覆盖测试")
    document.add_paragraph("摘要")
    document.add_paragraph("正文内容。")
    document.save(path)
    return path


def _write_zip(path: Path, members: dict[str, bytes]) -> Path:
    with ZipFile(path, "w", ZIP_DEFLATED) as archive:
        for name, value in members.items():
            archive.writestr(name, value)
    return path


def _core_types(*, overrides: tuple[str, ...] = ()) -> bytes:
    override_text = "".join(
        f'<Override PartName="/{part}" ContentType="application/xml"/>'
        for part in overrides
    )
    return (
        f'<Types xmlns="{CT_NS}">'
        '<Default Extension="xml" ContentType="application/xml"/>'
        f"{override_text}</Types>"
    ).encode()


def _document_xml(*, root: str = "document", body: str = "<w:body/>") -> bytes:
    return (f'<w:{root} xmlns:w="{W_NS}">{body}</w:{root}>').encode()


def _relationship_xml(*relations: str) -> bytes:
    return (
        f'<Relationships xmlns="{REL_NS}">{"".join(relations)}</Relationships>'.encode()
    )


def test_list_rules_prints_one_tabular_record_per_rule(capsys) -> None:
    assert cli.main(["list-rules"]) == 0

    lines = capsys.readouterr().out.splitlines()
    assert len(lines) == 49
    fields = lines[0].split("\t")
    assert fields[0] == "ER-MS-TITLE-001"
    assert len(fields) == 4
    assert lines[-1].startswith("ER-REF-CONTENT-FOOTNOTE-001\t")


def test_validate_rules_plain_text_reports_success(capsys) -> None:
    assert cli.main(["validate-rules"]) == 0

    assert capsys.readouterr().out == "OK: 49 rules validated\n"


def test_validate_rules_plain_text_reports_structured_errors(
    monkeypatch, capsys
) -> None:
    monkeypatch.setattr(cli, "validate_rules", lambda: ["manuscript.yaml::bad"])

    assert cli.main(["validate-rules"]) == 1

    assert capsys.readouterr().out == "ERROR manuscript.yaml::bad\n"


def test_validate_rules_json_returns_one_for_invalid_result(
    monkeypatch, capsys
) -> None:
    monkeypatch.setattr(
        cli,
        "validate_rules_structured",
        lambda: {
            "valid": False,
            "errors": [{"code": "file_missing"}],
            "rule_count": 0,
            "conflict_count": 0,
            "unresolved_count": 0,
            "source_count": 0,
            "error_count": 1,
        },
    )

    assert cli.main(["validate-rules", "--json"]) == 1

    payload = json.loads(capsys.readouterr().out)
    assert payload["valid"] is False
    assert payload["errors"][0]["code"] == "file_missing"


def test_inspect_rejects_an_output_directory_without_traceback(
    tmp_path: Path, capsys
) -> None:
    input_path = _minimal_docx(tmp_path / "paper.docx")
    output_dir = tmp_path / "already-a-directory"
    output_dir.mkdir()

    assert cli.main(["inspect", str(input_path), "--output", str(output_dir)]) == 2

    captured = capsys.readouterr()
    assert "Traceback" not in captured.err
    assert "ERROR DOCX:" in captured.err
    assert output_dir.is_dir()


def test_inspect_missing_input_with_existing_output_keeps_output_and_handles_samefile_error(
    tmp_path: Path, capsys
) -> None:
    missing = tmp_path / "missing.docx"
    output = tmp_path / "inspection.json"
    output.write_text("sentinel", encoding="utf-8")

    assert cli.main(["inspect", str(missing), "--output", str(output)]) == 2

    assert output.read_text(encoding="utf-8") == "sentinel"
    assert "Traceback" not in capsys.readouterr().err


def test_inspect_debug_re_raises_input_error(tmp_path: Path) -> None:
    missing = tmp_path / "missing.docx"
    output = tmp_path / "inspection.json"

    with pytest.raises(DocxInspectionError, match="does not exist"):
        cli.main(["inspect", str(missing), "--output", str(output), "--debug"])

    assert not output.exists()


def test_inspect_rejects_exact_input_output_path_without_modifying_docx(
    tmp_path: Path, capsys
) -> None:
    input_path = _minimal_docx(tmp_path / "paper.docx")
    before = input_path.read_bytes()

    assert cli.main(["inspect", str(input_path), "--output", str(input_path)]) == 2

    assert input_path.read_bytes() == before
    assert "same file" in capsys.readouterr().err.lower()


def test_lint_include_text_and_metadata_persists_opted_in_fields(
    tmp_path: Path,
) -> None:
    input_path = _minimal_docx(tmp_path / "metadata.docx")
    document = Document(input_path)
    document.core_properties.author = "CLI Author"
    document.core_properties.title = "CLI Title"
    document.save(input_path)
    output_dir = tmp_path / "audit"

    assert (
        cli.main(
            [
                "lint",
                str(input_path),
                "--output-dir",
                str(output_dir),
                "--include-text",
                "--include-metadata",
                "--exit-zero",
            ]
        )
        == 0
    )

    payload = json.loads((output_dir / "inspection.json").read_text(encoding="utf-8"))
    assert payload["paragraphs"][0]["text"] == "题目：CLI 覆盖测试"
    assert payload["core_properties"]["creator"] == "CLI Author"
    assert payload["core_properties"]["title"] == "CLI Title"


def test_redaction_drops_full_text_and_bounds_nested_previews() -> None:
    payload = {
        "full_text": "private document text",
        "nested": [
            {
                "text": "private paragraph",
                "text_preview": "x" * 120,
                "full_text": "private nested text",
            }
        ],
    }

    redacted = cli._without_full_text(payload)

    assert "full_text" not in redacted
    assert "full_text" not in redacted["nested"][0]
    assert "text" not in redacted["nested"][0]
    assert redacted["nested"][0]["text_preview"] == "x" * 80


def test_package_open_reports_missing_input(tmp_path: Path) -> None:
    path = tmp_path / "missing.docx"

    with pytest.raises(DocxInspectionError) as raised:
        DocxPackage.open(path)

    assert raised.value.kind == "missing_input"
    assert raised.value.path == path


def test_package_open_rejects_unsafe_member_names(tmp_path: Path) -> None:
    path = _write_zip(
        tmp_path / "unsafe.docx",
        {
            "[Content_Types].xml": _core_types(),
            "word/document.xml": _document_xml(),
            "../escape": b"unsafe",
        },
    )

    with pytest.raises(DocxInspectionError, match="unsafe member path") as raised:
        DocxPackage.open(path)

    assert raised.value.kind == "zip_member"


def test_package_open_rejects_duplicate_member_names(tmp_path: Path) -> None:
    path = tmp_path / "duplicate.docx"
    with warnings.catch_warnings():
        warnings.simplefilter("ignore", UserWarning)
        with ZipFile(path, "w", ZIP_DEFLATED) as archive:
            archive.writestr("[Content_Types].xml", _core_types())
            archive.writestr("word/document.xml", _document_xml())
            archive.writestr("word/document.xml", _document_xml())

    with pytest.raises(DocxInspectionError, match="duplicate member name") as raised:
        DocxPackage.open(path)

    assert raised.value.kind == "zip_member"


def test_package_open_enforces_member_count_limit(tmp_path: Path, monkeypatch) -> None:
    path = _write_zip(
        tmp_path / "many-members.docx",
        {
            "[Content_Types].xml": _core_types(),
            "word/document.xml": _document_xml(),
            "extra.bin": b"extra",
        },
    )
    monkeypatch.setattr(package_module, "MAX_ZIP_MEMBERS", 2)

    with pytest.raises(DocxInspectionError, match="too many members") as raised:
        DocxPackage.open(path)

    assert raised.value.kind == "resource_limit"


def test_package_open_enforces_member_size_limit(tmp_path: Path, monkeypatch) -> None:
    path = _write_zip(
        tmp_path / "large-member.docx",
        {
            "[Content_Types].xml": _core_types(),
            "word/document.xml": _document_xml(),
            "extra.bin": b"x" * 4096,
        },
    )
    monkeypatch.setattr(package_module, "MAX_MEMBER_UNCOMPRESSED_BYTES", 1024)

    with pytest.raises(DocxInspectionError, match="member exceeds") as raised:
        DocxPackage.open(path)

    assert raised.value.kind == "resource_limit"


def test_package_open_enforces_total_uncompressed_size_limit(
    tmp_path: Path, monkeypatch
) -> None:
    members = {
        "[Content_Types].xml": _core_types(),
        "word/document.xml": _document_xml(),
        "one.bin": b"1" * 128,
        "two.bin": b"2" * 128,
    }
    path = _write_zip(tmp_path / "large-total.docx", members)
    sizes = [info.file_size for info in ZipFile(path).infolist()]
    monkeypatch.setattr(package_module, "MAX_MEMBER_UNCOMPRESSED_BYTES", max(sizes) + 1)
    monkeypatch.setattr(package_module, "MAX_TOTAL_UNCOMPRESSED_BYTES", sum(sizes) - 1)

    with pytest.raises(DocxInspectionError, match="total uncompressed size") as raised:
        DocxPackage.open(path)

    assert raised.value.kind == "resource_limit"


def test_package_open_reports_corrupt_xml_part(tmp_path: Path) -> None:
    path = _write_zip(
        tmp_path / "bad-xml.docx",
        {
            "[Content_Types].xml": _core_types(),
            "word/document.xml": b"<w:document",
        },
    )

    with pytest.raises(DocxInspectionError, match="word/document.xml") as raised:
        DocxPackage.open(path)

    assert raised.value.kind == "xml_part"


def test_package_open_rejects_a_non_word_document_root(tmp_path: Path) -> None:
    path = _write_zip(
        tmp_path / "not-word.docx",
        {
            "[Content_Types].xml": _core_types(),
            "word/document.xml": _document_xml(root="notDocument", body=""),
        },
    )

    with pytest.raises(DocxInspectionError, match="not a Word document") as raised:
        DocxPackage.open(path)

    assert raised.value.kind == "core_part"


def test_package_open_rejects_declared_missing_part(tmp_path: Path) -> None:
    path = _write_zip(
        tmp_path / "missing-declared-part.docx",
        {
            "[Content_Types].xml": _core_types(overrides=("word/missing.xml",)),
            "word/document.xml": _document_xml(),
        },
    )

    with pytest.raises(DocxInspectionError, match="declares missing part") as raised:
        DocxPackage.open(path)

    assert raised.value.kind == "missing_part"


def test_package_relationship_adapters_cover_external_local_and_empty_targets(
    tmp_path: Path,
) -> None:
    relation_type = "http://example.invalid/relationship"
    relations = (
        f'<Relationship Id="rId1" Type="{relation_type}" '
        'Target="https://example.invalid" TargetMode="External"/>'
        f'<Relationship Id="rId2" Type="{relation_type}" Target="media.bin"/>'
        f'<Relationship Id="rId3" Type="{relation_type}" Target=""/>'
        f'<Relationship Id="rId4" Target="media.bin"/>'
    )
    path = _write_zip(
        tmp_path / "relationships.docx",
        {
            "[Content_Types].xml": _core_types(),
            "word/document.xml": _document_xml(),
            "word/media.bin": b"media",
            "word/_rels/document.xml.rels": _relationship_xml(relations),
        },
    )
    package = DocxPackage.open(path)

    relationships = package.relationships()
    assert set(relationships) == {"rId1", "rId2"}
    assert (
        package.relationship_target("word/document.xml", "rId1")[1]
        == "https://example.invalid"
    )
    assert (
        package.relationship_target("word/document.xml", "rId2")[1] == "word/media.bin"
    )
    assert package.relationship_target("word/document.xml", "missing") is None
    assert (
        package.resolve_target("word/document.xml", "/word/media.bin")
        == "word/media.bin"
    )
    assert package.resolve_target("word/document.xml", "media.bin") == "word/media.bin"


def test_package_read_xml_and_content_type_adapters_are_deterministic(
    tmp_path: Path,
) -> None:
    path = _write_zip(
        tmp_path / "accessors.docx",
        {
            "[Content_Types].xml": _core_types(),
            "word/document.xml": _document_xml(),
            "extra.bin": b"payload",
        },
    )
    package = DocxPackage.open(path)

    assert package.read("extra.bin") == b"payload"
    assert package.read("missing.bin") is None
    assert package.xml("word/document.xml") is package.document_root
    assert package.xml("missing.xml") is None
    assert package.content_type("word/document.xml") == "application/xml"
    assert package.content_type("word/other.xml") == "application/xml"
    assert package.content_type("word/other.bin") is None
    assert package.relationships() == {}

    package.parts["extra.xml"] = b"<broken"
    with pytest.raises(DocxInspectionError, match="extra.xml") as raised:
        package.xml("extra.xml")
    assert raised.value.kind == "xml_part"

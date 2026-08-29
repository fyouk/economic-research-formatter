from __future__ import annotations

import json
import os
from pathlib import Path

from docx import Document

from economic_research_formatter.cli import _without_full_text, main


def _minimal_docx(path: Path) -> Path:
    document = Document()
    document.add_paragraph("题目：合成测试稿")
    document.add_paragraph("摘要")
    document.add_paragraph("这是合成测试摘要。")
    document.add_paragraph("关键词：测试")
    document.add_paragraph("一、引言")
    document.add_paragraph("正文。")
    document.add_paragraph("长段落" * 40)
    document.add_paragraph("参考文献")
    document.add_paragraph("张三，2020，《测试》。")
    document.save(path)
    return path


def test_validate_rules_json_is_machine_readable(capsys):
    exit_code = main(["validate-rules", "--json"])

    payload = json.loads(capsys.readouterr().out)
    assert exit_code == 0
    assert payload["valid"] is True
    assert payload["rule_count"] == 49
    assert payload["errors"] == []


def test_inspect_cli_writes_deterministic_json(tmp_path):
    input_path = _minimal_docx(tmp_path / "中文论文.docx")
    first = tmp_path / "first.json"
    second = tmp_path / "second.json"

    assert main(["inspect", str(input_path), "--output", str(first)]) == 0
    assert main(["inspect", str(input_path), "--output", str(second)]) == 0

    first_payload = json.loads(first.read_text(encoding="utf-8"))
    assert first.read_bytes() == second.read_bytes()
    assert first_payload["schema_version"] == "1.0"
    assert first_payload["input"]["filename"] == "中文论文.docx"
    assert "text" not in first_payload["paragraphs"][0]
    assert len(first_payload["paragraphs"][0]["text_preview"]) <= 80


def test_inspect_cli_only_writes_full_text_when_explicitly_requested(tmp_path):
    input_path = _minimal_docx(tmp_path / "paper.docx")
    output = tmp_path / "full.json"

    assert main(
        ["inspect", str(input_path), "--output", str(output), "--include-text"]
    ) == 0

    payload = json.loads(output.read_text(encoding="utf-8"))
    assert payload["paragraphs"][0]["text"] == "题目：合成测试稿"
    long_paragraph = next(
        paragraph for paragraph in payload["paragraphs"] if len(paragraph.get("text", "")) > 80
    )
    assert len(long_paragraph["text_preview"]) == 80


def test_core_metadata_requires_a_separate_explicit_cli_opt_in(tmp_path):
    input_path = _minimal_docx(tmp_path / "metadata.docx")
    document = Document(input_path)
    document.core_properties.author = "Private Creator"
    document.core_properties.last_modified_by = "Private Modifier"
    document.core_properties.title = "Private Title"
    document.save(input_path)

    compact = tmp_path / "compact.json"
    text_only = tmp_path / "text-only.json"
    expanded = tmp_path / "metadata.json"
    assert main(["inspect", str(input_path), "--output", str(compact)]) == 0
    assert main(
        ["inspect", str(input_path), "--output", str(text_only), "--include-text"]
    ) == 0
    assert main(
        [
            "inspect",
            str(input_path),
            "--output",
            str(expanded),
            "--include-metadata",
        ]
    ) == 0

    for output in (compact, text_only):
        core = json.loads(output.read_text(encoding="utf-8"))["core_properties"]
        assert core["creator"]["present"] is True
        assert "Private Creator" not in json.dumps(core)
        assert "title" not in core
    expanded_core = json.loads(expanded.read_text(encoding="utf-8"))["core_properties"]
    assert expanded_core["creator"] == "Private Creator"
    assert expanded_core["last_modified_by"] == "Private Modifier"
    assert expanded_core["title"] == "Private Title"


def test_lint_cli_writes_all_reports_and_exit_zero_override(tmp_path):
    input_path = _minimal_docx(tmp_path / "paper.docx")
    output_dir = tmp_path / "audit"

    exit_code = main(
        ["lint", str(input_path), "--output-dir", str(output_dir), "--exit-zero"]
    )

    assert exit_code == 0
    assert (output_dir / "inspection.json").is_file()
    assert (output_dir / "audit.json").is_file()
    assert (output_dir / "audit.md").is_file()
    inspection = json.loads((output_dir / "inspection.json").read_text(encoding="utf-8"))
    assert all("text" not in paragraph for paragraph in inspection["paragraphs"])
    assert max(len(paragraph["text_preview"]) for paragraph in inspection["paragraphs"]) <= 80
    audit = json.loads((output_dir / "audit.json").read_text(encoding="utf-8"))
    previews = [
        finding["target"]["text_preview"]
        for finding in audit["findings"]
        if "text_preview" in finding["target"]
    ]
    assert previews and max(map(len, previews)) <= 80
    assert audit["schema_version"] == "1.0"
    assert audit["summary"]["by_status"]["ERROR"] >= 1
    assert "# 《经济研究》格式审计报告" in (
        output_dir / "audit.md"
    ).read_text(encoding="utf-8")


def test_lint_cli_returns_one_when_errors_exist(tmp_path):
    input_path = _minimal_docx(tmp_path / "paper.docx")

    assert main(["lint", str(input_path), "--output-dir", str(tmp_path / "audit")]) == 1


def test_cli_reports_invalid_docx_without_traceback(tmp_path, capsys):
    bad_input = tmp_path / "not-a-docx.docx"
    bad_input.write_text("not a zip package", encoding="utf-8")

    exit_code = main(["inspect", str(bad_input), "--output", str(tmp_path / "x.json")])

    captured = capsys.readouterr()
    assert exit_code == 2
    assert "Traceback" not in captured.err
    assert "DOCX" in captured.err


def test_inspect_rejects_hardlink_output_without_modifying_input(tmp_path, capsys):
    input_path = _minimal_docx(tmp_path / "paper.docx")
    before = input_path.read_bytes()
    output = tmp_path / "inspection.json"
    os.link(input_path, output)

    exit_code = main(["inspect", str(input_path), "--output", str(output)])

    assert exit_code == 2
    assert input_path.read_bytes() == before
    assert output.read_bytes() == before
    assert "same file" in capsys.readouterr().err.lower()


def test_lint_rejects_hardlink_report_without_modifying_input(tmp_path, capsys):
    input_path = _minimal_docx(tmp_path / "paper.docx")
    before = input_path.read_bytes()
    output_dir = tmp_path / "audit"
    output_dir.mkdir()
    os.link(input_path, output_dir / "inspection.json")

    exit_code = main(
        ["lint", str(input_path), "--output-dir", str(output_dir), "--exit-zero"]
    )

    assert exit_code == 2
    assert input_path.read_bytes() == before
    assert (output_dir / "inspection.json").read_bytes() == before
    assert not (output_dir / "audit.json").exists()
    assert "same file" in capsys.readouterr().err.lower()


def test_redaction_removes_content_text_but_preserves_numbering_text():
    payload = {
        "text": "private full paragraph",
        "text_preview": "x" * 120,
        "deleted_text": "private deleted text",
        "deleted_text_preview": "y" * 120,
        "numPr": {"text": "1.", "format": "decimal"},
    }

    redacted = _without_full_text(payload)

    assert "text" not in redacted
    assert "deleted_text" not in redacted
    assert redacted["text_preview"] == "x" * 80
    assert redacted["deleted_text_preview"] == "y" * 80
    assert redacted["numPr"]["text"] == "1."

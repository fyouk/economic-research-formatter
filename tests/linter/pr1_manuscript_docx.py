"""Real DOCX fixtures for PR1 manuscript-consumer tests.

The helpers deliberately exercise the public ``inspect_docx`` boundary.  They
only create temporary synthetic documents; no private manuscript is read or
written.
"""

from __future__ import annotations

from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from tests.docx_factory import make_synthetic_docx


def _set_latin_font(run, name: str) -> None:
    run.font.name = name
    rpr = run._r.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), name)


def make_table_latin_docx(tmp_path: Path) -> Path:
    """Create a body-plus-table document whose only Latin text is in a table."""

    document = Document()
    body = document.add_paragraph("正文")
    body_run = body.runs[0]
    _set_latin_font(body_run, "宋体")
    body_run.font.size = Pt(10.5)
    table = document.add_table(rows=1, cols=1)
    cell_run = table.cell(0, 0).paragraphs[0].add_run("Revenue 123")
    _set_latin_font(cell_run, "Arial")
    cell_run.font.size = Pt(9)
    path = tmp_path / "table-latin.docx"
    document.save(path)
    return path


def make_mixed_table_docx(tmp_path: Path) -> Path:
    """Create a table with mixed effective sizes for a pending relation contract."""

    document = Document()
    body = document.add_paragraph("正文基准")
    body_run = body.runs[0]
    _set_latin_font(body_run, "宋体")
    body_run.font.size = Pt(10.5)
    table = document.add_table(rows=1, cols=1)
    paragraph = table.cell(0, 0).paragraphs[0]
    first = paragraph.add_run("A")
    first.font.size = Pt(9)
    _set_latin_font(first, "Times New Roman")
    second = paragraph.add_run("B")
    second.font.size = Pt(10.5)
    _set_latin_font(second, "Times New Roman")
    path = tmp_path / "table-mixed.docx"
    document.save(path)
    return path


def make_valid_table_relation_docx(tmp_path: Path) -> Path:
    """Create a table/note pair whose Inspector comparisons are determinate."""

    document = Document()
    body = document.add_paragraph("正文基准")
    body_run = body.runs[0]
    _set_latin_font(body_run, "宋体")
    body_run.font.size = Pt(10.5)
    table = document.add_table(rows=1, cols=1)
    table_run = table.cell(0, 0).paragraphs[0].add_run("表格数据")
    _set_latin_font(table_run, "仿宋")
    table_run.font.size = Pt(9)
    note = document.add_paragraph("注：表格说明。")
    note_run = note.runs[0]
    _set_latin_font(note_run, "宋体")
    note_run.font.size = Pt(7.5)
    path = tmp_path / "table-valid-relations.docx"
    document.save(path)
    return path


def make_mathtype_docx(tmp_path: Path, *, with_omml: bool = False) -> Path:
    """Create a synthetic OLE document whose ProgID proves MathType."""

    path = make_synthetic_docx(
        tmp_path,
        filename="mathtype.docx",
        with_real_footnote=False,
        with_metadata_objects=True,
    )
    with ZipFile(path, "r") as source:
        members = {name: source.read(name) for name in source.namelist()}
    xml = members["word/document.xml"].replace(b'ProgID="Package"', b'ProgID="MathType 7.0"')
    members["word/document.xml"] = xml
    if not with_omml:
        # The base fixture has OMML paragraphs.  Remove only the equation
        # elements so this remains an OLE-only editor fixture.
        from lxml import etree

        root = etree.fromstring(xml)
        math_ns = "http://schemas.openxmlformats.org/officeDocument/2006/math"
        for node in root.xpath(".//m:oMath | .//m:oMathPara", namespaces={"m": math_ns}):
            parent = node.getparent()
            if parent is not None:
                parent.remove(node)
        members["word/document.xml"] = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    with ZipFile(path, "w", ZIP_DEFLATED) as target:
        for name, data in members.items():
            target.writestr(name, data)
    return path


__all__ = [
    "make_mathtype_docx",
    "make_mixed_table_docx",
    "make_table_latin_docx",
    "make_valid_table_relation_docx",
]

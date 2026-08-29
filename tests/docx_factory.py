"""Small, in-memory DOCX fixtures used by Inspector tests.

The factory intentionally creates the OOXML features that ``python-docx`` does
not expose as a first-class API (fields, OMML and a footnote part) after the
normal document has been saved.  It never reads or writes the private paper
fixture.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw
from lxml import etree


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
O_NS = "urn:schemas-microsoft-com:office:office"
OFFICE_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def _set_run_font(run, *, east_asia: str | None = None, ascii_font: str | None = None,
                  size: float | None = None, size_cs: float | None = None,
                  bold: bool | None = None, italic: bool | None = None) -> None:
    if east_asia or ascii_font:
        rpr = run._r.get_or_add_rPr()
        fonts = rpr.rFonts
        if fonts is None:
            fonts = OxmlElement("w:rFonts")
            rpr.insert(0, fonts)
        if east_asia:
            fonts.set(qn("w:eastAsia"), east_asia)
        if ascii_font:
            fonts.set(qn("w:ascii"), ascii_font)
            fonts.set(qn("w:hAnsi"), ascii_font)
    if size is not None:
        run.font.size = Pt(size)
    if size_cs is not None:
        rpr = run._r.get_or_add_rPr()
        sz_cs = rpr.find(qn("w:szCs"))
        if sz_cs is None:
            sz_cs = OxmlElement("w:szCs")
            rpr.append(sz_cs)
        sz_cs.set(qn("w:val"), str(round(size_cs * 2)))
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _add_field(paragraph, instruction: str, result: str = "") -> None:
    """Append a standard complex Word field to ``paragraph``."""
    begin = OxmlElement("w:r")
    begin_char = OxmlElement("w:fldChar")
    begin_char.set(qn("w:fldCharType"), "begin")
    begin.append(begin_char)
    instr_run = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    instr_run.append(instr)
    separate = OxmlElement("w:r")
    separate_char = OxmlElement("w:fldChar")
    separate_char.set(qn("w:fldCharType"), "separate")
    separate.append(separate_char)
    if result:
        result_run = OxmlElement("w:r")
        result_text = OxmlElement("w:t")
        result_text.text = result
        result_run.append(result_text)
        paragraph._p.append(result_run)
    end = OxmlElement("w:r")
    end_char = OxmlElement("w:fldChar")
    end_char.set(qn("w:fldCharType"), "end")
    end.append(end_char)
    for element in (begin, instr_run, separate, end):
        paragraph._p.append(element)


def _append_omml(paragraph, *, text: str = "x = y + 1", in_para: bool = False) -> None:
    omath = parse_xml(
        f'<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        f'<m:r><m:t>{text}</m:t></m:r></m:oMath>'
    )
    if in_para:
        wrapper = OxmlElement("m:oMathPara")
        wrapper.append(omath)
        paragraph._p.append(wrapper)
    else:
        paragraph._p.append(omath)


def _add_numpr(paragraph, *, num_id: int = 1, ilvl: int = 0) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    numid_el = OxmlElement("w:numId")
    numid_el.set(qn("w:val"), str(num_id))
    numpr.extend((ilvl_el, numid_el))


def _add_bookmark(paragraph, name: str = "fixture_bookmark", bookmark_id: str = "1") -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bookmark_id)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bookmark_id)
    paragraph._p.insert(1, start)
    paragraph._p.append(end)


def _add_hyperlink(paragraph, text: str = "fixture link", url: str = "https://example.invalid") -> None:
    document_part = paragraph.part
    relationship_id = document_part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_pr.append(color)
    run.append(run_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _png_bytes(mode: str) -> bytes:
    if mode == "transparent_bw":
        image = Image.new("RGBA", (32, 24), (255, 255, 255, 0))
        draw = ImageDraw.Draw(image)
        draw.line((3, 3, 28, 20), fill=(0, 0, 0, 255), width=2)
        draw.rectangle((8, 7, 20, 16), outline=(0, 0, 0, 255), width=1)
    elif mode == "bw":
        image = Image.new("RGB", (40, 30), "white")
        draw = ImageDraw.Draw(image)
        draw.ellipse((5, 5, 30, 24), outline="black", width=2)
    else:
        image = Image.new("RGB", (48, 36), "white")
        draw = ImageDraw.Draw(image)
        draw.rectangle((4, 4, 24, 29), fill=(200, 25, 25))
        draw.rectangle((24, 8, 43, 32), fill=(30, 100, 210))
    stream = BytesIO()
    image.save(stream, format="PNG")
    return stream.getvalue()


def _with_ooxml_parts(path: Path, *, footnotes: bool = True, comments: bool = True,
                      tracked_changes: bool = True, embedded: bool = True) -> None:
    """Add OOXML-only parts while preserving the normal python-docx package."""
    with ZipFile(path, "r") as source:
        members = {name: source.read(name) for name in source.namelist()}

    rels = parse_xml(members["word/_rels/document.xml.rels"])
    content_types = parse_xml(members["[Content_Types].xml"])
    document = parse_xml(members["word/document.xml"])
    rel_index = 10

    def add_rel(rel_type: str, target: str) -> str:
        nonlocal rel_index
        existing = [int(el.get("Id", "rId0")[3:]) for el in rels if el.get("Id", "").startswith("rId") and el.get("Id")[3:].isdigit()]
        rel_index = max(existing + [rel_index]) + 1
        rid = f"rId{rel_index}"
        rel = etree.Element(f"{{{REL_NS}}}Relationship")
        rel.set("Id", rid)
        rel.set("Type", rel_type)
        rel.set("Target", target)
        rels.append(rel)
        return rid

    def add_default_or_override(kind: str, name: str, content_type: str) -> None:
        if kind == "default":
            element = etree.Element(f"{{{CT_NS}}}Default")
            element.set("Extension", name)
        else:
            element = etree.Element(f"{{{CT_NS}}}Override")
            element.set("PartName", name)
        element.set("ContentType", content_type)
        content_types.append(element)

    # Namespace prefixes used by OxmlElement are registered in python-docx;
    # lxml still accepts these explicit Clark-notation tags.
    if footnotes:
        add_rel(f"{OFFICE_REL_NS}/footnotes", "footnotes.xml")
        add_default_or_override("override", "/word/footnotes.xml", "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml")
        members["word/footnotes.xml"] = (
            f'<w:footnotes xmlns:w="{W_NS}">'
            '<w:footnote w:type="separator" w:id="-1"><w:p/></w:footnote>'
            '<w:footnote w:type="continuationSeparator" w:id="0"><w:p/></w:footnote>'
            '<w:footnote w:id="1"><w:p><w:r><w:t>真实脚注</w:t></w:r></w:p></w:footnote>'
            '</w:footnotes>'
        ).encode()
        p = document.find(f".//{{{W_NS}}}body/{{{W_NS}}}p")
        if p is not None:
            ref_run = OxmlElement("w:r")
            ref = OxmlElement("w:footnoteReference")
            ref.set(qn("w:id"), "1")
            ref_run.append(ref)
            p.append(ref_run)

    if comments:
        comment_rel_type = f"{OFFICE_REL_NS}/comments"
        add_rel(comment_rel_type, "comments.xml")
        add_default_or_override("override", "/word/comments.xml", "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml")
        members["word/comments.xml"] = (
            f'<w:comments xmlns:w="{W_NS}" xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006">'
            '<w:comment w:id="0" w:author="fixture" w:date="2026-01-01T00:00:00Z"><w:p><w:r><w:t>comment</w:t></w:r></w:p></w:comment>'
            '</w:comments>'
        ).encode()
        p = document.find(f".//{{{W_NS}}}body/{{{W_NS}}}p")
        if p is not None:
            start = OxmlElement("w:commentRangeStart")
            start.set(qn("w:id"), "0")
            end = OxmlElement("w:commentRangeEnd")
            end.set(qn("w:id"), "0")
            p.insert(0, start)
            p.append(end)

    if tracked_changes:
        p = document.find(f".//{{{W_NS}}}body/{{{W_NS}}}p")
        if p is not None:
            ins = OxmlElement("w:ins")
            ins.set(qn("w:id"), "1")
            ins.set(qn("w:author"), "fixture")
            r = OxmlElement("w:r")
            t = OxmlElement("w:t")
            t.text = " inserted"
            r.append(t)
            ins.append(r)
            p.append(ins)
            deletion = OxmlElement("w:del")
            deletion.set(qn("w:id"), "2")
            deletion.set(qn("w:author"), "fixture")
            r = OxmlElement("w:r")
            t = OxmlElement("w:delText")
            t.text = " deleted"
            r.append(t)
            deletion.append(r)
            p.append(deletion)

    if embedded:
        embedded_rel_type = f"{OFFICE_REL_NS}/oleObject"
        ole_rel_id = add_rel(embedded_rel_type, "embeddings/oleObject1.bin")
        add_default_or_override("default", "bin", "application/vnd.openxmlformats-officedocument.oleObject")
        members["word/embeddings/oleObject1.bin"] = b"synthetic embedded object"
        p = document.find(f".//{{{W_NS}}}body/{{{W_NS}}}p")
        if p is not None:
            run = OxmlElement("w:r")
            obj = OxmlElement("w:object")
            ole = etree.Element(f"{{{O_NS}}}OLEObject")
            ole.set("Type", "Embed")
            ole.set("ProgID", "Package")
            ole.set(qn("r:id"), ole_rel_id)
            obj.append(ole)
            run.append(obj)
            p.append(run)

    members["word/_rels/document.xml.rels"] = etree.tostring(
        rels, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    members["[Content_Types].xml"] = etree.tostring(
        content_types, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    members["word/document.xml"] = etree.tostring(
        document, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    with ZipFile(path, "w", ZIP_DEFLATED) as target:
        for name, data in members.items():
            target.writestr(name, data.getvalue() if hasattr(data, "getvalue") else data)


def make_synthetic_docx(tmp_path: Path, filename: str = "synthetic.docx", *,
                        with_real_footnote: bool = True, with_metadata_objects: bool = True) -> Path:
    """Create a deterministic small DOCX exercising the Inspector contract."""
    tmp_path = Path(tmp_path)
    path = tmp_path / filename
    doc = Document()
    doc.core_properties.title = "Inspector fixture"
    doc.core_properties.author = "Fixture Author"

    # Four sections, with one landscape section.
    for index in range(3):
        section = doc.add_section()
        section.start_type = 2  # new page
        if index == 1:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    title = doc.paragraphs[0]
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("题目：")
    _set_run_font(run, east_asia="宋体", size=16)
    run = title.add_run("混合字体标题")
    _set_run_font(run, east_asia="黑体", ascii_font="Arial", size=14)
    _add_bookmark(title, "title")

    author = doc.add_paragraph("张三和李四")
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(author.runs[0], east_asia="宋体", size=12)

    abstract_heading = doc.add_paragraph("摘要")
    abstract_heading.style = "Heading 1"
    abstract = doc.add_paragraph("本文摘要正文，包含一个 citation (Smith, 2020)。")
    _set_run_font(abstract.runs[0], east_asia="仿宋", size=10.5)
    keywords = doc.add_paragraph("关键词：人工智能；财务错报")
    _set_run_font(keywords.runs[0], east_asia="宋体", size=10.5)

    toc = doc.add_paragraph()
    _add_field(toc, 'TOC \\o "1-3" \\h', "第一章\t1")
    toc.add_run(" ")
    _add_field(toc, "PAGEREF _Toc123 \\h", "1")

    headings = ["第1章 绪论", "1.1 研究背景", "1.1.1 研究问题", "（1）研究设计"]
    for index, text in enumerate(headings):
        paragraph = doc.add_paragraph(text)
        paragraph.style = f"Heading {min(index + 1, 4)}"
        if index == 0:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_font(paragraph.runs[0], east_asia="宋体", size=14 - index)

    equation = doc.add_paragraph("公式：")
    _append_omml(equation, in_para=False)
    where = doc.add_paragraph("其中  x 表示解释变量")
    _append_omml(where, text="y = ax + b", in_para=True)

    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "表1"
    table.cell(0, 1).text = "值"
    table.cell(1, 0).text = "A"
    table.cell(1, 1).text = "10"
    note = doc.add_paragraph("注：表格数据为合成夹具。")
    _set_run_font(note.runs[0], east_asia="宋体", size=9)

    image_modes = ("transparent_bw", "bw", "color")
    for index, mode in enumerate(image_modes, start=1):
        picture = doc.add_paragraph()
        image_path = tmp_path / f"fixture-{index}.png"
        image_path.write_bytes(_png_bytes(mode))
        picture.add_run().add_picture(str(image_path), width=Inches(1.2))
        caption = doc.add_paragraph(f"图{index}.1 合成图题")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_font(caption.runs[0], east_asia="黑体", size=9)

    citation = doc.add_paragraph("Ravallian and Chen (2001)；王一与李二（2020）；Smith et al. (2019)")
    _add_hyperlink(citation)

    reference_heading = doc.add_paragraph("参考文献")
    reference_heading.style = "Heading 1"
    references = [
        "Smith, J. (2019). An English Article. Journal, 1, 15-25.",
        "王一，李二（2020）：中文文献。经济研究，1，1—10。",
        "Brown, A. et al. (2018). Another article. Publisher.",
    ]
    for reference in references:
        paragraph = doc.add_paragraph(reference)
        _add_numpr(paragraph, num_id=1)
        _set_run_font(paragraph.runs[0], east_asia="宋体", ascii_font="Calibri", size=7.5)

    doc.save(path)
    _with_ooxml_parts(path, footnotes=with_real_footnote, comments=with_metadata_objects,
                       tracked_changes=with_metadata_objects, embedded=with_metadata_objects)
    return path


# Friendly aliases make the fixture importable by downstream tests without
# coupling them to one particular test name.
build_synthetic_docx = make_synthetic_docx
create_synthetic_docx = make_synthetic_docx
make_test_docx = make_synthetic_docx


__all__ = [
    "build_synthetic_docx",
    "create_synthetic_docx",
    "make_synthetic_docx",
    "make_test_docx",
]

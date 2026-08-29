from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

from economic_research_formatter.docx.inspector import inspect_docx


def test_inspector_preserves_raw_and_effective_run_formatting_sources(tmp_path: Path) -> None:
    path = tmp_path / "styles.docx"
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"].font.size = Pt(12)
    paragraph = document.add_paragraph()
    paragraph.style = styles["Normal"]
    paragraph.add_run("继承")
    direct = paragraph.add_run("直排")
    direct.font.name = "黑体"
    direct.font.size = Pt(16)
    document.save(path)

    report = inspect_docx(path)
    runs = report["paragraphs"][-1]["runs"]
    inherited_format = runs[0]["formatting"]
    direct_format = runs[1]["formatting"]

    assert "raw" in inherited_format
    assert inherited_format["effective"]
    assert inherited_format["source"]
    assert direct_format["raw"]
    assert direct_format["effective"]
    assert direct_format["source"]["font"]["eastAsia"] in {"direct", "character_style", "paragraph_style", "docDefaults", "theme", "unknown"}
    assert direct_format["effective"]["size_pt"] == 16
    assert direct_format["effective"]["size_cs_pt"] == 11
    assert direct_format["effective"]["size_cs_pt"] != direct_format["effective"]["size_pt"]
